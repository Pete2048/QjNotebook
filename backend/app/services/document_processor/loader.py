from typing import List, Dict, Any, Optional
from pathlib import Path
import mimetypes
from langchain_community.document_loaders import (
    TextLoader,
    PyPDFLoader,
    Docx2txtLoader,
    CSVLoader,
    UnstructuredMarkdownLoader
)
from langchain_core.documents import Document
import base64
import io
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

class DocumentLoader:
    """文档加载器，支持多种文件格式"""
    
    def __init__(self):
        self.supported_extensions = {
            '.txt': self._load_text,
            '.md': self._load_markdown,
            '.pdf': self._load_pdf,
            '.docx': self._load_docx,
            '.doc': self._load_docx,
            '.csv': self._load_csv,
        }
    
    def load_from_path(self, file_path: str, metadata: Dict[str, Any] = None) -> List[Document]:
        """从文件路径加载文档"""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = path.suffix.lower()
        if extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {extension}")
        
        loader_func = self.supported_extensions[extension]
        documents = loader_func(str(path))
        
        # 添加元数据
        base_metadata = {
            "source": str(path),
            "filename": path.name,
            "file_type": extension,
            "file_size": path.stat().st_size,
        }
        
        if metadata:
            base_metadata.update(metadata)
        
        for doc in documents:
            doc.metadata.update(base_metadata)
        
        return documents
    
    def load_from_text(self, text: str, metadata: Dict[str, Any] = None) -> List[Document]:
        """从文本字符串加载文档"""
        base_metadata = {
            "source": "text_input",
            "file_type": "text",
            "content_length": len(text),
        }
        
        if metadata:
            base_metadata.update(metadata)
        
        return [Document(page_content=text, metadata=base_metadata)]
    
    def load_from_base64(self, base64_content: str, file_type: str, metadata: Dict[str, Any] = None) -> List[Document]:
        """从 base64 编码内容加载文档"""
        try:
            # 解码 base64
            content_bytes = base64.b64decode(base64_content)
            
            if file_type.lower() == 'pdf' or metadata.get('isPDF'):
                return self._load_pdf_from_bytes(content_bytes, metadata)
            elif file_type.lower() in ['docx', 'doc']:
                return self._load_docx_from_bytes(content_bytes, metadata)
            else:
                # 尝试作为文本处理
                text = content_bytes.decode('utf-8')
                return self.load_from_text(text, metadata)
        
        except Exception as e:
            raise ValueError(f"Failed to load document from base64: {str(e)}")
    
    def _load_text(self, file_path: str) -> List[Document]:
        """加载文本文件"""
        loader = TextLoader(file_path, encoding='utf-8')
        return loader.load()
    
    def _load_markdown(self, file_path: str) -> List[Document]:
        """加载 Markdown 文件"""
        try:
            loader = UnstructuredMarkdownLoader(file_path)
            return loader.load()
        except Exception:
            # 如果 unstructured 失败，回退到文本加载
            return self._load_text(file_path)
    
    def _load_pdf(self, file_path: str) -> List[Document]:
        """加载 PDF 文件"""
        loader = PyPDFLoader(file_path)
        return loader.load()
    
    def _load_pdf_from_bytes(self, content_bytes: bytes, metadata: Dict[str, Any] = None) -> List[Document]:
        """从字节流加载 PDF"""
        try:
            # 使用 PyPDF2 直接处理字节流
            pdf_file = io.BytesIO(content_bytes)
            pdf_reader = PdfReader(pdf_file)
            
            documents = []
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text.strip():  # 只添加非空页面
                    page_metadata = {
                        "page": page_num + 1,
                        "total_pages": len(pdf_reader.pages),
                        **(metadata or {})
                    }
                    documents.append(Document(
                        page_content=text,
                        metadata=page_metadata
                    ))
            
            return documents
        
        except Exception as e:
            raise ValueError(f"Failed to load PDF from bytes: {str(e)}")
    
    def _load_docx(self, file_path: str) -> List[Document]:
        """加载 DOCX 文件"""
        loader = Docx2txtLoader(file_path)
        return loader.load()
    
    def _load_docx_from_bytes(self, content_bytes: bytes, metadata: Dict[str, Any] = None) -> List[Document]:
        """从字节流加载 DOCX"""
        try:
            docx_file = io.BytesIO(content_bytes)
            doc = DocxDocument(docx_file)
            
            # 提取所有段落文本
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            full_text = '\n'.join(text_content)
            
            base_metadata = {
                "file_type": "docx",
                "paragraphs": len(text_content),
                **(metadata or {})
            }
            
            return [Document(page_content=full_text, metadata=base_metadata)]
        
        except Exception as e:
            raise ValueError(f"Failed to load DOCX from bytes: {str(e)}")
    
    def _load_csv(self, file_path: str) -> List[Document]:
        """加载 CSV 文件"""
        loader = CSVLoader(file_path)
        return loader.load()
    
    def get_supported_extensions(self) -> List[str]:
        """获取支持的文件扩展名"""
        return list(self.supported_extensions.keys())
    
    def is_supported(self, file_path: str) -> bool:
        """检查文件是否支持"""
        extension = Path(file_path).suffix.lower()
        return extension in self.supported_extensions
    
    def detect_file_type(self, file_path: str) -> Optional[str]:
        """检测文件类型"""
        mime_type, _ = mimetypes.guess_type(file_path)
        extension = Path(file_path).suffix.lower()
        
        type_mapping = {
            '.txt': 'text/plain',
            '.md': 'text/markdown',
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword',
            '.csv': 'text/csv',
        }
        
        return type_mapping.get(extension, mime_type)